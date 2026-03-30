import io
import os
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析PDF文件"""
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            print(f"PDF parse error: {e}")
        return text
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """解析Word文件"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"DOCX parse error: {e}")
        return text
    
    @staticmethod
    def parse_txt(file_path: str) -> str:
        """解析文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"TXT parse error: {e}")
            return ""
    
    @staticmethod
    def parse(file_path: str) -> str:
        """根据文件类型自动选择解析方法"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return DocumentParser.parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentParser.parse_docx(file_path)
        elif ext in ['.txt', '.md']:
            return DocumentParser.parse_txt(file_path)
        else:
            return ""
    
    @staticmethod
    def parse_from_content(content: bytes, file_format: str) -> str:
        """从内存内容解析"""
        if file_format == 'pdf':
            try:
                reader = PdfReader(io.BytesIO(content))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except Exception as e:
                print(f"PDF parse error: {e}")
                return ""
        
        elif file_format in ['docx', 'doc']:
            try:
                doc = DocxDocument(io.BytesIO(content))
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
            except Exception as e:
                print(f"DOCX parse error: {e}")
                return ""
        
        elif file_format in ['txt', 'md']:
            try:
                return content.decode('utf-8')
            except Exception as e:
                print(f"TXT parse error: {e}")
                return ""
        
        return ""
